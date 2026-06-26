import os
from docx import Document as DocxDocument
from app.schemas.document import Document

class DOCXLoader:
    def load(self, path: str) -> Document:
        doc = DocxDocument(path)

        parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            if not text:
                continue

            style = paragraph.style.name

            if style.startswith("Heading"):
                parts.append(f"\n# {text}\n")
            else:
                parts.append(text)

        text = "\n".join(parts)

        return Document(
            source="docx",
            filename=os.path.basename(path),
            text=text
        )